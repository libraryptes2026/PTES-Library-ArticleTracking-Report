import streamlit as st
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import pandas as pd
from datetime import datetime
import io
import docx
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# --- 1. PAGE CONFIGURATION & STYLING ---
st.set_page_config(
    page_title="PTES Reading Articles Tracker",
    page_icon="📚",
    layout="wide"
)

# --- Stable Native Word Page Number Generator (Prevents File Corruption) ---
def add_robust_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

# --- Helper Function for Pristine Administrative Formatting ---
def format_cell_borders_and_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    shading_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
    tcPr.append(shading_elm)
    
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')          
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  
        tcBorders.append(border)
    tcPr.append(tcBorders)

# --- 2. SECURE DATABASE CONNECTION ---
def connect_to_sheets():
    try:
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds_info = st.secrets["gspread_creds"]
        creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_info, scope)
        client = gspread.authorize(creds)
        return client
    except Exception as e:
        st.error(f"Database Connection Error: {e}")
        return None

client = connect_to_sheets()

# --- 🧠 HIGH-PERFORMANCE CACHING FUNCTIONS ENGINE ---
@st.cache_data(ttl=600)  # Keeps data cached safely in memory for 10 minutes to prevent 429 Errors
def fetch_registry_rows(target_spreadsheet):
    """Securely downloads and caches student records."""
    # We create an isolated internal client execution layer inside the cache pool
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    creds_info = st.secrets["gspread_creds"]
    creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_info, scope)
    local_client = gspread.authorize(creds)
    sheet = local_client.open(target_spreadsheet).worksheet("Student_Registry")
    return sheet.get_all_values()

@st.cache_data(ttl=600)
def fetch_leaderboard_range(target_spreadsheet):
    """Securely fetches leaderboard ranges separately to prevent heavy sheet reloading."""
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    creds_info = st.secrets["gspread_creds"]
    creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_info, scope)
    local_client = gspread.authorize(creds)
    sheet = local_client.open(target_spreadsheet).worksheet("Student_Registry")
    return sheet.get("N225:P227")


# --- 3. SIDEBAR ---
with st.sidebar:
    st.image("https://upload.wikimedia.org/wikipedia/commons/thumb/c/c3/Flag_of_Brunei.svg/180px-Flag_of_Brunei.svg.png", width=100)
    st.title("PTES Library Services")
    st.markdown("### 📋 System Guidelines")
    
    with st.expander("📖 Reading Articles Rules"):
        st.write("""
        1. **Article Schedule:** Chief Librarian releases 24-25 articles monthly (Monday–Thursday & Saturday).
        2. **Honesty & Integrity:** Tallies must correspond strictly to verified student readings.
        3. **Librarian Control:** Only authorized library staff may unlock and submit records.
        """)
        
    st.divider()
    # Manual bypass fallback control in case they want fresh sheet values immediately
    if st.button("🔄 Force Clear Cached Memory", use_container_width=True):
        st.cache_data.clear()
        st.success("App cache cleared successfully!")
        st.rerun()
        
    st.divider()
    st.markdown("""
        <div style='font-size: 11px; font-weight: bold; text-align: center; margin-bottom: 8px;'>
            ⬜ Perseverance &nbsp; 🟥 Trustworthiness &nbsp; 🔵 Exemplary &nbsp; 🟡 Self-reliance &nbsp; 🟢
        </div>
        <div style='font-size: 11px; text-align: center; color: #888;'>
            Web Developer: Miss Hjh Nurul Haziqah HN (PTES Computer Science)
        </div>
    """, unsafe_allow_html=True)

# --- 4. MAIN INTERFACE ---
st.title("📚 PTES Reading Articles Tracker")
st.markdown("Automating the tracking of student reading habits cleanly and efficiently.")
st.markdown("---")

if client:
    admin_pass = st.text_input("🗝️ Enter Librarian Credentials to Unlock Portal:", type="password")
    
    if admin_pass == st.secrets["admin_password"]:
        st.success("🔓 Librarian Access Granted.")
        st.markdown("---")
        
        st.markdown("### 🗂️ Step 1: Select Cohort Level")
        cohort_choice = st.selectbox(
            "Choose Student Cohort Database to access:",
            ["Lower Sixth (BE Classes)", "Upper Sixth (AE Classes)"]
        )
        
        if cohort_choice == "Lower Sixth (BE Classes)":
            target_spreadsheet = "Articles_Tracker_DB_BE"
        else:
            target_spreadsheet = "Articles_Tracker_DB_AE"
            
        st.info(f"Connected to live database: **{target_spreadsheet}**")
        st.markdown("---")

        try:
            # Replaced heavy sheet processing with optimized memory lookup wrapper
            all_rows = fetch_registry_rows(target_spreadsheet)
            
            headers = all_rows[0]
            data_rows = all_rows[1:]
            df_registry = pd.DataFrame(data_rows, columns=headers)
            
            ordered_classes = []
            class_registry = {}
            for _, row in df_registry.iterrows():
                form_class = str(row.get("Form Class", "")).strip()
                student_name = str(row.get("Student Name", "")).strip()
                if form_class and student_name and not form_class.startswith("TOP") and form_class != "Form Class":
                    if form_class not in class_registry:
                        class_registry[form_class] = []
                        ordered_classes.append(form_class)
                    class_registry[form_class].append(student_name)
                    
        except Exception as e:
            st.error(f"Error accessing Google Sheet: {e}")
            df_registry = pd.DataFrame()
            class_registry = {}
            ordered_classes = []

        portal_tab1, portal_tab2, portal_tab3 = st.tabs([
            "📝 Submit Tally Logs", 
            "📊 Class Statistics Report", 
            "🏆 Top Readers Leaderboard"
        ])

        # ==================== TAB 1: LOG INSERTER ====================
        with portal_tab1:
            if not class_registry:
                st.warning("No registry data found. Please check your Google Sheet tabs.")
            else:
                st.markdown("### 📝 Enter Student Tally Details")
                col1, col2 = st.columns(2)
                
                with col1:
                    selected_class = st.selectbox("1. Select Form Class:", ordered_classes, key="challenge_class")
                    student_options = class_registry.get(selected_class, [])
                    selected_student = st.selectbox("2. Select Student Name:", student_options, key="challenge_student")
                
                with col2:
                    months_list = ["March", "April", "May", "June", "July", "August", "September", "October"]
                    selected_month = st.selectbox("3. Select Challenge Month:", months_list)
                    article_count = st.number_input("4. Amount of Reading Articles Done:", min_value=1, max_value=31, step=1)
                
                st.markdown("---")
                st.write(f"📂 **5. Enter the specific calendar dates for the {int(article_count)} materials read:**")
                
                month_map = {
                    "March": 3, "April": 4, "May": 5, "June": 6, 
                    "July": 7, "August": 8, "September": 9, "October": 10
                }
                target_month_num = month_map.get(selected_month, datetime.now().month)
                current_year = datetime.now().year
                default_calendar_date = datetime(current_year, target_month_num, 1)
                
                reading_dates = []
                cols = st.columns(4)
                
                for i in range(int(article_count)):
                    col_index = i % 4
                    with cols[col_index]:
                        date_val = st.date_input(
                            f"Article #{i+1} Date", 
                            value=default_calendar_date,
                            key=f"{selected_month}_challenge_date_{i}"
                        )
                        reading_dates.append(date_val.strftime("%d/%m/%Y"))
                
                st.markdown("---")
                student_remarks = st.text_input(
                    "📝 6. Additional Remarks / Student Status Notes (Optional):", 
                    value="", placeholder="e.g., Special consideration"
                )
                
                st.markdown("---")
                
                if st.button("🔥 Submit Tally Data Logs", use_container_width=True):
                    try:
                        challenge_db = client.open(target_spreadsheet).worksheet("Reading_Article_DB")
                        dates_string = ", ".join(reading_dates)
                        
                        new_tally_row = [
                            str(datetime.now().strftime("%d/%m/%Y %H:%M:%S")),
                            str(selected_class),
                            str(selected_student),
                            str(selected_month),
                            int(article_count),  
                            str(dates_string),
                            str(student_remarks).strip()
                        ]
                        
                        challenge_db.append_row(new_tally_row, value_input_option="USER_ENTERED")
                        
                        # ✨ CLEAR THE MEMORY CACHE IN REAL-TIME UPON ADDING NEW TALLIES
                        st.cache_data.clear()
                        
                        st.success(f"🎉 Success! Recorded {article_count} articles into Column E.")
                        st.balloons()
                    except Exception as e:
                        st.error(f"Failed to record data: {e}")

        # ==================== TAB 2: LIVE REPORT GENERATOR ====================
        with portal_tab2:
            st.markdown("### 📊 Class Reading Overview")
            
            if df_registry.empty:
                st.warning("Waiting for connection to active Google Sheet...")
            else:
                report_class = st.selectbox("Select Target Class to View:", ordered_classes)
                
                ordered_months = ["MARCH", "APRIL", "MAY", "JUNE", "JULY", "AUGUST", "SEPTEMBER", "OCTOBER"]
                current_month_idx = datetime.now().month  
                
                visible_columns = ["Form Class", "Student Name"]
                for m_name in ordered_months:
                    found_col = next((col for col in df_registry.columns if col.upper().startswith(m_name[:3])), None)
                    if found_col:
                        month_num = month_map.get(m_name.capitalize(), 12)
                        if month_num <= current_month_idx:
                            visible_columns.append(found_col)
                
                filtered_df = df_registry[df_registry["Form Class"] == report_class].copy()
                display_table = filtered_df[visible_columns].reset_index(drop=True)
                display_table = display_table.replace(["#N/A", "nan", ""], "0")
                
                month_cols_only = [c for c in visible_columns if c not in ["Form Class", "Student Name"]]
                for col in month_cols_only:
                    display_table[col] = pd.to_numeric(display_table[col], errors='coerce').fillna(0).astype(int)
                
                display_table["TOTAL"] = display_table[month_cols_only].sum(axis=1)
                display_table.insert(0, "NO.", range(1, len(display_table) + 1))
                
                st.dataframe(display_table, use_container_width=True, hide_index=True)
                
                # --- CLASS REPORT GENERATION ENGINE ---
                doc = Document()
                for section in doc.sections:
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = Inches(11.0)
                    section.page_height = Inches(8.5)
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)
                    
                    footer_p = section.footer.paragraphs[0]
                    footer_p.alignment = 2  
                    footer_run = footer_p.add_run("Page ")
                    footer_run.font.size = Pt(9)
                    footer_run.font.name = 'Arial'
                    add_robust_page_number(footer_run)
                
                title_p = doc.add_paragraph()
                title_run = title_p.add_run("PUSAT TINGKATAN ENAM SENGKURONG\nSTUDENT READING ARTICLES SUMMARY REPORT")
                title_run.bold = True
                title_run.font.size = Pt(12)
                title_run.font.name = 'Arial'
                title_p.alignment = 1
                
                meta_p = doc.add_paragraph()
                meta_run = meta_p.add_run(f"CLASS: {report_class}  |  DATABASE COHORT: {cohort_choice.upper()}\nREPORT GENERATED ON: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
                meta_run.font.size = Pt(9.5)
                meta_run.italic = True
                meta_p.alignment = 1
                
                doc.add_paragraph().paragraph_format.space_after = Pt(12)
                
                word_cols = list(display_table.columns)
                table = doc.add_table(rows=1, cols=len(word_cols))
                table.autofit = True
                
                hdr_cells = table.rows[0].cells
                for idx, col_name in enumerate(word_cols):
                    hdr_cells[idx].text = str(col_name).upper()
                    run = hdr_cells[idx].paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    format_cell_borders_and_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
                
                for _, row in display_table.iterrows():
                    row_cells = table.add_row().cells
                    for idx, col_name in enumerate(word_cols):
                        row_cells[idx].text = str(row[col_name])
                        run = row_cells[idx].paragraphs[0].runs[0]
                        run.font.size = Pt(9.5)
                        run.font.name = 'Arial'
                        if col_name in ["TOTAL", "NO."]:
                            run.font.bold = True
                        format_cell_borders_and_margins(row_cells[idx], top=100, bottom=100, left=150, right=150)
                
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                st.markdown("---")
                st.download_button(
                    label=f"📥 Download Official {report_class} Reading Summary (.docx)",
                    data=doc_io,
                    file_name=f"PTES_Reading_Report_{report_class}_{datetime.now().strftime('%b%Y')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

        # ==================== TAB 3: LEADERBOARD & EXPORT ====================
        with portal_tab3:
            st.markdown("### 🏆 Cohort Top 3 Honors Registry")
            try:
                # Replaced heavy sheet get range call with memory optimized lookup
                top_cells = fetch_leaderboard_range(target_spreadsheet)
                
                if top_cells and len(top_cells) > 0:
                    cols_dash = st.columns(len(top_cells))
                    medals = ["🥇 1st Place", "🥈 2nd Place", "🥉 3rd Place"]
                    
                    for idx, data_row in enumerate(top_cells):
                        if len(data_row) >= 3:
                            with cols_dash[idx]:
                                st.metric(label=f"{medals[idx]} ({data_row[1]})", value=f"{data_row[0]}", delta=f"{data_row[2]} Articles")
                    
                    leaderboard_data = []
                    for data_row in top_cells:
                        if len(data_row) >= 3:
                            leaderboard_data.append({
                                "Rank Placement": medals[len(leaderboard_data)],
                                "Form Class": data_row[1],
                                "Student Name": data_row[0],
                                "Total Articles Read": int(data_row[2]) if str(data_row[2]).isdigit() else data_row[2]
                            })
                    
                    df_leaderboard = pd.DataFrame(leaderboard_data)
                    st.markdown("---")
                    st.dataframe(df_leaderboard, use_container_width=True, hide_index=True)
                    
                    leader_doc = Document()
                    for section in leader_doc.sections:
                        section.orientation = WD_ORIENT.LANDSCAPE
                        section.page_width = Inches(11.0)
                        section.page_height = Inches(8.5)
                        section.top_margin = Inches(0.5)
                        section.bottom_margin = Inches(0.5)
                        section.left_margin = Inches(0.5)
                        section.right_margin = Inches(0.5)
                        
                        f_p = section.footer.paragraphs[0]
                        f_p.alignment = 2
                        f_run = f_p.add_run("Page ")
                        f_run.font.size = Pt(9)
                        add_robust_page_number(f_run)
                    
                    l_title_p = leader_doc.add_paragraph()
                    l_title_p.add_run("PUSAT TINGKATAN ENAM SENGKURONG\nTOP READERS LEADERBOARD RECOGNITION REPORT").bold = True
                    l_title_p.alignment = 1
                    
                    l_table = leader_doc.add_table(rows=1, cols=4)
                    l_table.autofit = True
                    
                    l_doc_io = io.BytesIO()
                    leader_doc.save(l_doc_io)
                    l_doc_io.seek(0)
                    
                    st.markdown("---")
                    st.download_button(
                        label="📥 Download Official Landscape Leaderboard Report (.docx)",
                        data=l_doc_io,
                        file_name=f"PTES_Leaderboard_{cohort_choice.replace(' ', '_')}_{datetime.now().strftime('%b%Y')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="leaderboard_download_btn"
                    )
            except Exception as e:
                st.error(f"Error executing leaderboard logic: {e}")

    elif admin_pass != "":
        st.error("❌ Invalid Credentials. Access Denied.")
